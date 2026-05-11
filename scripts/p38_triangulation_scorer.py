#!/usr/bin/env python3
"""P38 triangulation scorer — classify each visible HUDOC paragraph by role
using three independent signal sources:

  1. DOCX paragraph style + run formatting (paragraph.style.name,
     paragraph.runs[].bold/italic/font.size, alignment, indent)
  2. HTML CSS class properties (font-style, font-weight, font-size,
     text-align, text-indent, margin-top, page-break-before)
  3. Text + sequence (leading number, monotonicity, all-caps, position,
     sibling context)

Each role is scored independently from all three sources; the best role
wins.  Disagreements between sources are surfaced via a confidence band.

Roles:
  main_paragraph   — numbered ¶ in main judgment body
  heading          — section / sub-section title
  quote            — indented blockquote (Ju_Quot)
  operative        — operative-part list item ("Holds…", "Declares…")
  metadata         — cover, judges composition, court details
  signature        — registrar / president signature
  footer           — "Done in English …" notification
  separate_opinion — dissenting / concurring opinion paragraph
  table_cell       — annex / applicant-list table cell
  toc              — table of contents entry
  unknown          — couldn't classify confidently

Usage:
    python3 scripts/p38_triangulation_scorer.py --case-id 001-249494
    python3 scripts/p38_triangulation_scorer.py --case-list /tmp/sample50.txt
"""
from __future__ import annotations
import argparse
import json
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Length
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from bs4 import BeautifulSoup
from bs4.element import Tag

ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = Path.home() / "Desktop" / "HUDOC-Docx"
HTML_DIR = Path.home() / "Desktop" / "HUDOC-Html"

PARA_NUM_RE = re.compile(r"^\s*(\d+)\.\s+")
TOC_LINE_RE = re.compile(r"\t\d{1,4}\s*$")
OPERATIVE_VERBS_RE = re.compile(r"^\s*(Declares?|Holds?|Dismisses?|Decides?|Strikes?|Joins?|Reserves?|Rejects?)\b", re.I)
PROSE_VERB_RE = re.compile(r"\b(was|were|is|are|had|have|has|did|does|do|been|being|shall|will|may|should)\b", re.I)
HEADING_NEEDLES_EN = {
    "PROCEDURE", "THE FACTS", "THE LAW", "FOR THESE REASONS",
    "SUBJECT MATTER OF THE CASE", "APPLICATION OF ARTICLE",
    "JUST SATISFACTION", "OTHER COMPLAINTS",
    "ALLEGED VIOLATION OF", "THE COURT'S ASSESSMENT",
    "RELEVANT LEGAL FRAMEWORK", "RELEVANT DOMESTIC LAW",
    "THE CIRCUMSTANCES OF THE CASE",
}
HEADING_NEEDLES_FR = {
    "PROCÉDURE", "EN FAIT", "EN DROIT", "PAR CES MOTIFS",
    "SATISFACTION ÉQUITABLE", "APPLICATION DE L'ARTICLE",
    "OBJET DE L'AFFAIRE", "AUTRES VIOLATIONS ALLÉGUÉES",
    "VIOLATION ALLÉGUÉE DE", "APPRÉCIATION DE LA COUR",
    "LES CIRCONSTANCES DE L'ESPÈCE", "DROIT INTERNE PERTINENT",
}
SEPARATE_OPINION_TITLE_RE = re.compile(
    r"\b(SEPARATE|DISSENTING|CONCURRING|JOINT)\s+(PARTLY\s+)?"
    r"(CONCURRING|DISSENTING|JOINT|OPINION|OPINIONS|DECLARATION)\b",
    re.I,
)
FOOTER_RE = re.compile(r"^(Done in (English|French)|Fait en (anglais|fran[cç]ais))", re.I)

CSS_PROP_RE = re.compile(r"([a-z-]+)\s*:\s*([^;]+)\s*;")


# ---------------------------------------------------------------------------
# Block model
# ---------------------------------------------------------------------------


@dataclass
class BlockFeatures:
    """Everything we know about one paragraph-block, from all three sources."""
    # Universal
    text: str = ""
    text_norm: str = ""            # whitespace-collapsed lower
    position_idx: int = 0
    position_ratio: float = 0.0    # 0.0 = first, 1.0 = last
    leading_number: Optional[int] = None

    # DOCX-side
    docx_style: str = ""
    docx_align: str = ""
    docx_left_indent_pt: Optional[float] = None
    docx_first_line_indent_pt: Optional[float] = None
    docx_run_bold: bool = False
    docx_run_italic: bool = False
    docx_run_size_pt: Optional[float] = None

    # HTML-side (resolved from CSS class definitions)
    html_class: str = ""
    html_font_style: str = ""
    html_font_weight: str = ""
    html_font_size_pt: Optional[float] = None
    html_text_align: str = ""
    html_text_indent_pt: Optional[float] = None
    html_margin_top_pt: Optional[float] = None
    html_margin_bottom_pt: Optional[float] = None
    html_margin_left_pt: Optional[float] = None
    html_page_break_before: bool = False
    html_in_table: bool = False

    # Sequence
    prev_main_number: Optional[int] = None
    expected_next_main: Optional[int] = None

    def __post_init__(self):
        self.text_norm = re.sub(r"\s+", " ", (self.text or "").strip()).lower()


@dataclass
class Classification:
    role: str
    score: float
    confidence: str   # "high" / "medium" / "low"
    scores: dict      # per-role scores for debugging
    reasons: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# DOCX feature extractor
# ---------------------------------------------------------------------------


def _pt(value: Length | None) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value.pt)
    except Exception:
        return None


def _iter_visible(parent, in_table: bool = False):
    """Yield (Paragraph, in_table_flag) in body order, descending into tables.
    Mirrors P37's iter_visible_paragraphs so annex/applicant-list tables
    are not silently dropped."""
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent), in_table
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for tr in table._tbl.tr_lst:
                for tc in tr.tc_lst:
                    cell = _Cell(tc, table)
                    yield from _iter_visible(cell, True)


def _is_ooxml_zip(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            return f.read(2) == b"PK"
    except Exception:
        return False


def extract_docx_features(docx_path: Path) -> list[BlockFeatures]:
    """Return BlockFeatures-without-html for each visible paragraph in DOCX
    (including paragraphs nested inside table cells — these are typical of
    Pop-C committee judgments with applicant-list annexes).

    Legacy pre-OOXML binary .doc files are transparently converted to
    DOCX via macOS textutil before parsing (HUDOC occasionally serves
    such files via the /conversion/docx endpoint).
    """
    if not _is_ooxml_zip(docx_path):
        import subprocess, tempfile
        with tempfile.TemporaryDirectory(prefix="hudoc_legacy_") as tmp:
            tmp_dir = Path(tmp)
            tmp_docx = tmp_dir / "converted.docx"
            try:
                subprocess.run(
                    ["textutil", "-convert", "docx", "-output",
                     str(tmp_docx), str(docx_path)],
                    check=True, capture_output=True, timeout=60,
                )
            except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
                # Best-effort; give up and let python-docx raise
                pass
            if tmp_docx.exists():
                doc = Document(str(tmp_docx))
            else:
                doc = Document(str(docx_path))
    else:
        doc = Document(str(docx_path))
    blocks: list[BlockFeatures] = []
    for p, in_table in _iter_visible(doc):
        text = (p.text or "").strip()
        if not text:
            continue
        fmt = p.paragraph_format
        # First non-whitespace run's formatting is representative enough
        first_run = next((r for r in p.runs if r.text and r.text.strip()), None)
        b = BlockFeatures(
            text=text,
            docx_style=p.style.name if p.style else "",
            docx_align=str(fmt.alignment) if fmt.alignment is not None else "",
            docx_left_indent_pt=_pt(fmt.left_indent),
            docx_first_line_indent_pt=_pt(fmt.first_line_indent),
            docx_run_bold=bool(first_run and first_run.bold),
            docx_run_italic=bool(first_run and first_run.italic),
            docx_run_size_pt=_pt(first_run.font.size) if first_run and first_run.font.size else None,
        )
        b.html_in_table = in_table   # also picked up from HTML later if available
        m = PARA_NUM_RE.match(text)
        if m:
            b.leading_number = int(m.group(1))
        blocks.append(b)
    n = len(blocks) or 1
    for i, b in enumerate(blocks):
        b.position_idx = i
        b.position_ratio = i / max(1, n - 1)
    return blocks


# ---------------------------------------------------------------------------
# HTML feature merger
# ---------------------------------------------------------------------------


def parse_css_classes(html: str) -> dict[str, dict[str, str]]:
    """Extract `.sXXXXXXXX { prop: val; ... }` map from the inline <style>."""
    classes: dict[str, dict[str, str]] = {}
    m = re.search(r"<style[^>]*>(.*?)</style>", html, re.S | re.I)
    if not m:
        return classes
    style = m.group(1)
    for cls_match in re.finditer(r"\.([a-zA-Z][\w-]*)\s*\{([^}]*)\}", style):
        cls = cls_match.group(1)
        body = cls_match.group(2)
        props: dict[str, str] = {}
        for prop_match in re.finditer(r"([a-z-]+)\s*:\s*([^;]+)\s*;", body):
            props[prop_match.group(1).strip()] = prop_match.group(2).strip()
        # last `prop: value` without trailing `;`
        last = re.search(r"([a-z-]+)\s*:\s*([^;]+)\s*$", body)
        if last:
            props[last.group(1).strip()] = last.group(2).strip()
        classes[cls] = props
    return classes


def _pt_from_css(val: str | None) -> Optional[float]:
    if not val:
        return None
    m = re.match(r"\s*([0-9.]+)\s*pt\s*", val)
    if not m:
        return None
    try:
        return float(m.group(1))
    except ValueError:
        return None


def merge_html_features(blocks: list[BlockFeatures], html_path: Path) -> None:
    """Match DOCX blocks against HTML <p>/<div> nodes by normalised text,
    in source order.  Populate HTML-side fields in place."""
    if not html_path.exists():
        return
    html = html_path.read_text(encoding="utf-8", errors="ignore")
    classes = parse_css_classes(html)

    # Collect ordered text-bearing nodes from the HTML body.
    soup = BeautifulSoup(html, "html.parser")
    html_nodes: list[Tag] = []
    for tag in soup.find_all(["p", "div", "td", "th", "li"]):
        txt = re.sub(r"\s+", " ", tag.get_text(" ", strip=True))
        if txt:
            html_nodes.append(tag)

    # Walk both lists; match by normalised text equality (first occurrence).
    htm_i = 0
    for b in blocks:
        norm_b = re.sub(r"\s+", " ", (b.text or "").strip())
        # find next HTML node whose normalised text contains b.text (or vice versa)
        j = htm_i
        while j < len(html_nodes):
            norm_h = re.sub(r"\s+", " ", html_nodes[j].get_text(" ", strip=True))
            if norm_h.startswith(norm_b[:60]) or norm_b.startswith(norm_h[:60]):
                tag = html_nodes[j]
                b.html_class = tag.get("class", [""])[0] if tag.get("class") else ""
                props: dict[str, str] = {}
                # combine all classes' properties
                for cls in tag.get("class", []):
                    props.update(classes.get(cls, {}))
                b.html_font_style = props.get("font-style", "")
                b.html_font_weight = props.get("font-weight", "")
                b.html_font_size_pt = _pt_from_css(props.get("font-size"))
                b.html_text_align = props.get("text-align", "")
                b.html_text_indent_pt = _pt_from_css(props.get("text-indent"))
                b.html_margin_top_pt = _pt_from_css(props.get("margin-top"))
                b.html_margin_bottom_pt = _pt_from_css(props.get("margin-bottom"))
                b.html_margin_left_pt = _pt_from_css(props.get("margin-left"))
                b.html_page_break_before = props.get("page-break-before") == "always"
                b.html_in_table = tag.name in ("td", "th") or bool(tag.find_parent("table"))
                htm_i = j + 1
                break
            j += 1


def populate_sequence(blocks: list[BlockFeatures]) -> None:
    last_main: Optional[int] = None
    max_seen = 0
    accept_count = 0
    for b in blocks:
        b.prev_main_number = last_main
        b.expected_next_main = (last_main + 1) if last_main else 1
        n = b.leading_number
        if n is not None:
            if accept_count > 0 and n <= max_seen:
                continue  # quote / non-main; don't update
            last_main = n
            max_seen = max(max_seen, n)
            accept_count += 1


# ---------------------------------------------------------------------------
# Per-role scorers
# ---------------------------------------------------------------------------


def looks_like_heading_only(text: str) -> bool:
    """All-caps section header without inline numbered ¶."""
    if not text or len(text) > 220:
        return False
    if re.search(r"\b\d+\.\s+[A-Z]", text):
        return False
    has_lower = any(c.islower() for c in text)
    has_upper = any(c.isupper() for c in text)
    return has_upper and not has_lower


def text_matches_heading_needle(text: str) -> bool:
    upper = text.upper()
    return any(needle in upper for needle in (HEADING_NEEDLES_EN | HEADING_NEEDLES_FR))


def score_main_paragraph(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style in ("Ju_Para", "Ju_Para_Last", "Ju_Para Char"):
        s += 5; why.append("docx:Ju_Para")
    if f.html_text_indent_pt and f.html_text_indent_pt > 5:
        s += 2; why.append("html:body-indent")
    if f.html_text_align == "justify":
        s += 1; why.append("html:justify")
    if f.html_font_style == "italic":
        s -= 3; why.append("html:italic (quote-like)")
    if f.html_font_weight in ("bold", "700"):
        s -= 2; why.append("html:bold (heading-like)")
    if f.html_in_table:
        s -= 4; why.append("html:in-table")
    if f.leading_number is not None:
        if f.expected_next_main and f.leading_number == f.expected_next_main:
            s += 5; why.append(f"seq:matches-next ({f.leading_number})")
        elif f.prev_main_number is None and f.leading_number <= 3:
            s += 5; why.append("seq:starts-at-1-2-3")
        elif f.prev_main_number and f.leading_number > f.prev_main_number:
            s += 1; why.append(f"seq:above-prev (skip {f.prev_main_number}→{f.leading_number})")
        elif f.prev_main_number and f.leading_number <= f.prev_main_number:
            s -= 5; why.append(f"seq:back-jump ({f.prev_main_number}→{f.leading_number})")
    else:
        s -= 2; why.append("no leading number")
    return s, why


def score_quote(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style == "Ju_Quot":
        s += 6; why.append("docx:Ju_Quot")
    if f.html_font_style == "italic":
        s += 4; why.append("html:italic")
    if f.html_margin_left_pt and f.html_margin_left_pt > 12:
        s += 3; why.append("html:left-margin")
    if f.html_font_size_pt and f.html_font_size_pt <= 10:
        s += 1; why.append("html:small-font")
    if f.text.startswith(("“", '"', "«")):
        s += 2; why.append("text:opens-quote")
    if f.leading_number is not None and f.prev_main_number and f.leading_number <= f.prev_main_number:
        s += 3; why.append("seq:non-monotonic")
    return s, why


def score_heading(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style.startswith("Ju_H_") or f.docx_style == "Ju_H_Head":
        s += 6; why.append(f"docx:{f.docx_style}")
    if f.html_font_weight in ("bold", "700"):
        s += 3; why.append("html:bold")
    if f.html_text_align == "center":
        s += 2; why.append("html:centered")
    if f.html_font_size_pt and f.html_font_size_pt >= 14:
        s += 3; why.append(f"html:font-{f.html_font_size_pt}")
    if f.html_margin_top_pt and f.html_margin_top_pt >= 24:
        s += 2; why.append("html:big-margin-top")
    if looks_like_heading_only(f.text):
        s += 4; why.append("text:all-caps")
    if text_matches_heading_needle(f.text):
        s += 3; why.append("text:heading-needle")
    if re.match(r"^([A-Z]\.|\([a-z]\)|[IVX]+\.)\s+[A-Z]", f.text) and len(f.text) <= 100:
        s += 2; why.append("text:sub-heading-marker")
    if f.leading_number is not None and re.match(r"^\s*\d+\.\s+[A-Z][a-z]", f.text):
        if len(f.text) < 80:
            s += 1; why.append("text:short-numbered (sub-heading?)")
        else:
            s -= 4; why.append("text:long-numbered-prose")
    if PROSE_VERB_RE.search(f.text) and len(f.text) > 100:
        s -= 3; why.append("text:prose-verb")
    return s, why


def score_operative(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style.startswith("Ju_List"):
        s += 6; why.append(f"docx:{f.docx_style}")
    if OPERATIVE_VERBS_RE.match(f.text):
        s += 5; why.append("text:operative-verb")
    if f.html_text_indent_pt is not None and f.html_text_indent_pt < 0:
        s += 2; why.append("html:hanging-indent")
    return s, why


def score_metadata(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style in ("Dec_H_Title", "Dec_H_Case", "Ju_Title", "Ju_Case",
                         "Ju_Judges", "Ju_Court", "ECHR_Cover_Title_4",
                         "ECHR_Placeholder"):
        s += 6; why.append(f"docx:{f.docx_style}")
    if f.position_ratio < 0.1 and f.html_text_align == "center":
        s += 3; why.append("html:top-centered")
    if f.html_font_size_pt and f.html_font_size_pt >= 16:
        s += 2; why.append("html:large-font")
    if re.match(r"^[A-Z\s.,()]+$", f.text) and 4 <= len(f.text) <= 60 and f.position_ratio < 0.1:
        s += 2; why.append("text:cover-shape")
    return s, why


def score_signature(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style == "Ju_Signed":
        s += 6; why.append("docx:Ju_Signed")
    if "Registrar" in f.text or "President" in f.text or "Greffier" in f.text:
        s += 3; why.append("text:registrar/president")
    if f.position_ratio > 0.9:
        s += 2; why.append("position:near-end")
    return s, why


def score_footer(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if FOOTER_RE.match(f.text):
        s += 6; why.append("text:done-in-english")
    if f.docx_style == "Ju_Para_Last":
        s += 2; why.append("docx:Ju_Para_Last")
    if f.position_ratio > 0.85:
        s += 1; why.append("position:near-end")
    return s, why


def score_separate_opinion(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style.startswith("Opi_"):
        s += 6; why.append(f"docx:{f.docx_style}")
    if SEPARATE_OPINION_TITLE_RE.search(f.text):
        s += 4; why.append("text:opinion-title")
    if f.position_ratio > 0.7:
        s += 1; why.append("position:after-operative")
    return s, why


def score_table_cell(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.html_in_table:
        s += 5; why.append("html:in-table")
    return s, why


def score_toc(f: BlockFeatures) -> tuple[float, list[str]]:
    s, why = 0.0, []
    if f.docx_style.lower().startswith("toc"):
        s += 6; why.append(f"docx:{f.docx_style}")
    if TOC_LINE_RE.search(f.text):
        s += 5; why.append("text:tab-page-no")
    return s, why


ROLE_SCORERS = {
    "main_paragraph":   score_main_paragraph,
    "quote":            score_quote,
    "heading":          score_heading,
    "operative":        score_operative,
    "metadata":         score_metadata,
    "signature":        score_signature,
    "footer":           score_footer,
    "separate_opinion": score_separate_opinion,
    "table_cell":       score_table_cell,
    "toc":              score_toc,
}


def classify_block(f: BlockFeatures) -> Classification:
    scores: dict[str, float] = {}
    reasons: dict[str, list[str]] = {}
    for role, scorer in ROLE_SCORERS.items():
        s, why = scorer(f)
        scores[role] = s
        reasons[role] = why
    role, top_score = max(scores.items(), key=lambda kv: kv[1])
    second = sorted(scores.values(), reverse=True)[1] if len(scores) > 1 else 0.0
    margin = top_score - second
    if top_score < 2:
        confidence = "low"
        role = "unknown"
    elif top_score >= 6 and margin >= 2:
        confidence = "high"
    elif top_score >= 4:
        confidence = "medium"
    else:
        confidence = "low"
    return Classification(
        role=role,
        score=top_score,
        confidence=confidence,
        scores=scores,
        reasons=reasons.get(role, []),
    )


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------


def classify_case(cid: str, verbose: bool = False) -> dict:
    docx_path = DOCX_DIR / f"{cid}.docx"
    html_path = HTML_DIR / f"{cid}.html"
    if not docx_path.exists():
        return {"cid": cid, "error": "docx missing"}
    try:
        blocks = extract_docx_features(docx_path)
    except Exception as e:
        return {"cid": cid, "error": f"docx parse: {type(e).__name__}: {str(e)[:80]}"}
    if html_path.exists():
        merge_html_features(blocks, html_path)
        html_used = True
    else:
        html_used = False
    populate_sequence(blocks)
    out_blocks = []
    role_counts: dict[str, int] = {}
    confidence_counts: dict[str, int] = {}
    for b in blocks:
        c = classify_block(b)
        role_counts[c.role] = role_counts.get(c.role, 0) + 1
        confidence_counts[c.confidence] = confidence_counts.get(c.confidence, 0) + 1
        if verbose:
            out_blocks.append({
                "text": b.text[:80],
                "role": c.role,
                "confidence": c.confidence,
                "score": round(c.score, 2),
                "docx_style": b.docx_style,
                "html_class": b.html_class,
                "leading_n": b.leading_number,
                "reasons": c.reasons,
            })
    return {
        "cid": cid,
        "html_used": html_used,
        "total_blocks": len(blocks),
        "role_counts": role_counts,
        "confidence_counts": confidence_counts,
        "blocks": out_blocks if verbose else [],
    }


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--case-id", action="append", default=[])
    ap.add_argument("--case-list")
    ap.add_argument("--verbose", action="store_true")
    ap.add_argument("--limit", type=int, default=0)
    args = ap.parse_args()

    cids: list[str] = list(args.case_id)
    if args.case_list:
        cids += [l.strip() for l in Path(args.case_list).read_text().splitlines() if l.strip()]
    if args.limit and len(cids) > args.limit:
        cids = cids[: args.limit]
    if not cids:
        print("usage: --case-id <id> [--case-id <id>] | --case-list <file>", file=sys.stderr)
        return 2

    for cid in cids:
        result = classify_case(cid, verbose=args.verbose)
        if args.verbose:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            cnt = result.get("role_counts", {})
            conf = result.get("confidence_counts", {})
            cnt_s = ", ".join(f"{k}={v}" for k, v in sorted(cnt.items()))
            conf_s = ", ".join(f"{k}={v}" for k, v in sorted(conf.items()))
            html = "html" if result.get("html_used") else "no-html"
            print(f"{cid}  blocks={result.get('total_blocks',0):>4}  [{html}]  {cnt_s}  conf:[{conf_s}]")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
